#🔷****
#🔷****
#🔷****
#🔷****
### استيراد المكتبات المطلوبة لتشغيل تطبيق Flask والتعامل مع البيانات، الملفات، الصور، كلمات المرور، Excel/Word، Google Drive وغيرها
from datetime import timedelta, datetime
from flask import Flask, render_template, request, jsonify, send_from_directory, send_file, redirect, url_for, session, flash
from flask_socketio import SocketIO, emit  # ✅ مكتبة الشات المباشر
import pandas as pd
from dateutil import parser
import os
import sqlite3
import json
import shutil
import re
import io
import hashlib
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

#🔷****
import json
from datetime import datetime, timedelta

HISTORY_FILE = 'chat_history.json'

def load_history():
    try:
        with open(HISTORY_FILE, 'r') as f:
            return json.load(f)
    except:
        return []

def save_message(user, content):
    history = load_history()
    history.append({
        'user': user,
        'message': content,
        'timestamp': datetime.now().isoformat()
    })
    with open(HISTORY_FILE, 'w') as f:
        json.dump(history, f)

#🔷****
#🔷****
#🔷****
### تهيئة تطبيق Flask وضبط المفتاح السري للجلسات

app = Flask(__name__)
app.secret_key = 'your-very-secret-key'

# ⬇️ أضف هذا السطر بعد تهيئة Flask مباشرة
socketio = SocketIO(app)

@app.context_processor
def inject_session_data():
    return {
        'session_username': session.get('username', ''),
        'session_company': session.get('company', ''),
        'session_facility': session.get('facility', ''),
        'session_login_time': session.get('login_time', '')
    }
#🔷****
#🔷****
#🔷****
#🔷****


#🔷****
#🔷****
#🔷****
#🔷****
### مسار تعديل بيانات المستخدم – يقوم فقط بتحويل المستخدم إلى صفحة إدارة المستخدمين

#🔷****
#🔷****
#🔷****
#🔷****
### عرض صفحة إدارة المستخدمين – تعرض قائمة المستخدمين، بيانات المستخدم المحدد، صلاحياته، وسجلات نشاطه

### ✅ 7 - Route: إدارة المستخدمين - صفحة manage_users
@app.route('/manage_users', methods=['GET', 'POST'])
def manage_users(): 
    print("🔍 SESSION DEBUG:", dict(session))  # ✅ طباعة بيانات الجلسة لتشخيص السبب

    if 'user_id' not in session:
        return redirect(url_for('login'))

    # ✅ التحقق من صلاحية الوصول بناءً على الدور أو اسم المستخدم Ar
    if session.get('username', '').strip().lower() != 'ar' and session.get('role') != 'data_admin':
        flash("⛔ ليس لديك صلاحية الوصول إلى صفحة إدارة المستخدمين", "danger")
        return redirect(url_for('index'))

    conn = sqlite3.connect("data.db")
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    c.execute("SELECT username FROM users WHERE is_active = 1")
    all_users = [row['username'] for row in c.fetchall()]

    selected_username = request.args.get('username')

    user_data = {}
    permissions = []
    user_logs = []

    if selected_username:
        c.execute("SELECT * FROM users WHERE username = ?", (selected_username,))
        user_data = dict(c.fetchone() or {})

        c.execute("SELECT permission FROM user_permissions WHERE user_id = (SELECT id FROM users WHERE username = ?)", (selected_username,))
        permissions = [row['permission'] for row in c.fetchall()]

        c.execute("""
            SELECT timestamp || ' - ' || action AS log_entry
            FROM activity_logs
            WHERE username = ?
            ORDER BY timestamp DESC
            LIMIT 20
        """, (selected_username,))
        user_logs = [row['log_entry'] for row in c.fetchall()]

    c.execute("SELECT DISTINCT company FROM users")
    companies = [row['company'] for row in c.fetchall() if row['company']]

    c.execute("SELECT DISTINCT facility FROM users")
    facilities = [row['facility'] for row in c.fetchall() if row['facility']]

    conn.close()

    return render_template(
        'manage_users.html',
        all_users=all_users,
        selected_username=selected_username,
        username=selected_username,
        first_name=user_data.get('first_name', ''),
        last_name=user_data.get('last_name', ''),
        email=user_data.get('email', ''),
        company=user_data.get('company', ''),
        facility=user_data.get('facility', ''),
        user_role=user_data.get('role', ''),
        is_active=user_data.get('is_active', 1),
        last_login=user_data.get('last_login', ''),
        last_modified=user_data.get('last_modified', ''),
        permissions=permissions,
        user_logs=user_logs,
        companies=companies,
        facilities=facilities,
        role=session.get('role', '')
    )

#🔷****
#🔷****
#🔷****
#🔷****
### ✅ 8 - Route: حفظ وتحديث بيانات وصلاحيات المستخدم - /add_permissions
@app.route('/add_permissions', methods=['POST'])
def add_permissions():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    username = request.form.get('target_username')
    new_password = request.form.get('new_password')
    role = request.form.get('role') or 'viewer'
    is_active = int(request.form.get('is_active') or 1)

    first_name = request.form.get('first_name', '')
    last_name = request.form.get('last_name', '')
    email = request.form.get('email', '')
    company = request.form.get('company', '')
    facility = request.form.get('facility', '')

    # ✅ دعم إضافة شركة أو منشأة جديدة من الحقول اليدوية
    if request.form.get('new_company'):
        company = request.form.get('new_company')
    if request.form.get('new_facility'):
        facility = request.form.get('new_facility')

    # ✅ كل الصلاحيات الممكنة (التي تظهر في صفحة الإدارة)
    all_possible_perms = [
        'can_upload_pm',
        'can_upload_asset',
        'can_delete_excel',
        'can_generate_reports',
        'can_edit_records',
        'can_view_zip',
        'can_add_users',
        'can_view_edit_asset',
        'can_view_edit'
    ]

    # ✅ استخراج الصلاحيات المحددة من الـ form
    granted_perms = request.form.getlist("permissions") or [
        perm for perm in all_possible_perms if request.form.get(perm)
    ]

    # ✅ ربط بقاعدة البيانات
    conn = sqlite3.connect('data.db')
    c = conn.cursor()

    # ✅ تحديث بيانات المستخدم
    if new_password:
        hashed_pw = generate_password_hash(new_password)
        c.execute('''
            UPDATE users
            SET password=?, role=?, is_active=?,
                first_name=?, last_name=?, email=?,
                company=?, facility=?, last_modified=CURRENT_TIMESTAMP
            WHERE username=?
        ''', (hashed_pw, role, is_active, first_name, last_name, email, company, facility, username))
    else:
        c.execute('''
            UPDATE users
            SET role=?, is_active=?,
                first_name=?, last_name=?, email=?,
                company=?, facility=?, last_modified=CURRENT_TIMESTAMP
            WHERE username=?
        ''', (role, is_active, first_name, last_name, email, company, facility, username))

    # ✅ حذف الصلاحيات القديمة
    c.execute("DELETE FROM user_permissions WHERE user_id = (SELECT id FROM users WHERE username = ?)", (username,))

    # ✅ إضافة الصلاحيات الجديدة (بعد إزالة بادئة can_)
    for perm in granted_perms:
        clean_perm = perm.split('_', 1)[1] if perm.startswith("can_") else perm
        c.execute("INSERT INTO user_permissions (user_id, permission) VALUES ((SELECT id FROM users WHERE username = ?), ?)", (username, clean_perm))

    conn.commit()
    conn.close()

    flash("✅ Permissions and user data updated successfully.", "success")
    return redirect(url_for('manage_users', username=username))

#🔷****
@app.route('/update_permissions', methods=['POST'])
def update_permissions():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('data.db')
    c = conn.cursor()

    # الصلاحيات التي تظهر في الصفحة
    all_permissions = [
        'upload_pm', 'upload_asset', 'delete_excel',
        'generate_reports', 'edit_records', 'view_zip',
        'add_users', 'view_edit_asset', 'view_edit'
    ]

    # جلب أسماء المستخدمين من الفورم
    usernames = request.form.getlist('usernames')

    for username in usernames:
        # الحصول على user_id من اسم المستخدم
        user_row = c.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
        if not user_row:
            continue
        user_id = user_row[0]

        # حذف الصلاحيات القديمة
        c.execute("DELETE FROM user_permissions WHERE user_id = ?", (user_id,))

        # إضافة الصلاحيات الجديدة من الفورم
        for perm in all_permissions:
            checkbox_name = f"{username}_{perm}"
            if checkbox_name in request.form:
                c.execute("INSERT INTO user_permissions (user_id, permission) VALUES (?, ?)", (user_id, perm))

    conn.commit()
    conn.close()
    flash("✅ تم حفظ التعديلات بنجاح", "success")
    return redirect(url_for('permissions_admin'))

#🔷****
#🔷****
#🔷****
### ✅ 9 - Route: تعديل سجل معين من ملف Excel - صفحة /edit (مع تتبع Debug)
@app.route('/edit', methods=['GET'])
def edit():
    if 'user_id' not in session:
        print("🚫 User not logged in")
        return redirect(url_for('login'))

    if session.get('username') != 'Ar' and not has_permission(session['user_id'], 'edit_records'):
        print("🚫 User does not have 'edit_records' permission")
        return "⛘️ You do not have permission to access this page", 403

    file_type = request.args.get('type', '')
    sheet_name = request.args.get('sheet', '')
    row_index = int(request.args.get('row', -1))

    print(f"📥 Edit request received: type={file_type}, sheet={sheet_name}, row={row_index}")

    if row_index < 0 or not sheet_name:
        print("⛔ Invalid request: missing row or sheet")
        return "⛔ Invalid request", 400

    # ✅ Load file from database instead of disk
    result = get_uploaded_excel(file_type)
    if not result:
        print(f"⛔ No uploaded file found for type: {file_type}")
        return f"⛔ No uploaded file for type: {file_type}", 404

    _, file_data, _ = result
    wb = load_workbook(io.BytesIO(file_data))
    print(f"📂 Available sheets: {wb.sheetnames}")

    if sheet_name not in wb.sheetnames:
        print(f"⛔ Sheet '{sheet_name}' not found in workbook")
        return f"⛔ Sheet not found: {sheet_name}", 404

    sheet = wb[sheet_name]
    headers = [cell.value for cell in sheet[1]]

    try:
        row_data = [cell.value for cell in sheet[row_index + 1]]
    except IndexError:
        print(f"⛔ Row {row_index} not found in sheet {sheet_name}")
        return "⛔ Requested row not found in file", 404

    data = dict(zip(headers, row_data))
    data["sheet"] = sheet_name
    data["row"] = row_index
    data["file_type"] = file_type

    print("✅ Data prepared for sending to template")
    # (Here you would typically return render_template with data)
    return render_template('edit_page.html', data=data)

#🔷****
#🔷****
#🔷****
#🔷****

### تسجيل مستخدم جديد مع إدخال البيانات الكاملة + حفظه في قاعدة البيانات + إعطاؤه صلاحيات كاملة

@app.route('/register', methods=['POST'])
def register():
    first_name = request.form.get('first_name')
    last_name = request.form.get('last_name')
    email = request.form.get('email')
    username = request.form.get('username')
    password = request.form.get('password')
    company = request.form.get('company')
    facility = request.form.get('facility')

    if company == '__new__':
        company = request.form.get('custom_company') or 'Unknown'
    if facility == '__new__':
        facility = request.form.get('custom_facility') or 'Unknown'

    hashed_password = generate_password_hash(password)

    try:
        with sqlite3.connect('data.db') as conn:
            c = conn.cursor()

            # تحقق من أن اسم المستخدم غير مستخدم مسبقًا
            c.execute("SELECT id FROM users WHERE username = ?", (username,))
            if c.fetchone():
                flash("❌ هذا المستخدم موجود بالفعل.", "danger")
                return redirect(url_for('index'))

            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

            # إضافة المستخدم ببيانات أساسية وصلاحيات 0
            c.execute('''
                INSERT INTO users (
                    first_name, last_name, email, username, password, company, facility,
                    role, is_active, last_login, last_modified,
                    can_upload_pm, can_upload_asset, can_delete_excel,
                    can_generate_reports, can_edit_records, can_view_zip, can_add_users
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                first_name, last_name, email, username, hashed_password, company, facility,
                'admin', 1, now, now,
                1, 1, 1, 1, 1, 1, 1
            ))

            # جلب معرف المستخدم المُنشأ
            user_id = c.lastrowid

            # إعطاء كل الصلاحيات
            full_permissions = [
                'upload_pm', 'upload_asset', 'delete_excel', 'generate_reports',
                'view_edit', 'view_edit_asset', 'view_dashboard', 'can_add_users'
            ]
            for perm in full_permissions:
                c.execute("INSERT INTO user_permissions (user_id, permission) VALUES (?, ?)", (user_id, perm))

            # تحميل الصلاحيات
            c.execute("SELECT permission FROM user_permissions WHERE user_id = ?", (user_id,))
            permissions = [row[0] for row in c.fetchall()]

        # تخزين بيانات الجلسة
        session.clear()
        session.permanent = True
        app.permanent_session_lifetime = timedelta(hours=6)
        session['user_id']     = user_id
        session['username']    = username.strip()
        session['company']     = company
        session['facility']    = facility
        session['permissions'] = permissions
        session['role']        = 'admin'
        session['login_time']  = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        session.modified = True

        print("✅ SESSION SET AFTER REGISTER:", dict(session))
        flash("✅ تم إنشاء الحساب وتسجيل الدخول بنجاح!", "success")
        return redirect(url_for('index'))

    except Exception as e:
        print("ERROR in register:", e)
        flash("❌ حدث خطأ أثناء إنشاء الحساب.", "danger")
        return redirect(url_for('register'))

### ✅ - Login Route: معالجة تسجيل الدخول وتخزين الصلاحيات
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        conn = sqlite3.connect("data.db")
        c = conn.cursor()

        # ✅ جلب بيانات المستخدم مع الدور
        c.execute("SELECT id, password, company, facility, role FROM users WHERE username = ?", (username,))
        user = c.fetchone()

        if user and check_password_hash(user[1], password):
            user_id  = user[0]
            company  = user[2]
            facility = user[3]
            role     = user[4] or ''

            # ✅ تحميل الصلاحيات
            c.execute("SELECT permission FROM user_permissions WHERE user_id = ?", (user_id,))
            permissions = [row[0] for row in c.fetchall()]
            conn.close()

            # ✅ حفظ بيانات الجلسة
            session.clear()
            session.permanent = True
            app.permanent_session_lifetime = timedelta(hours=6)

            session['user_id']     = user_id
            session['username']    = username.strip()
            session['company']     = company
            session['facility']    = facility
            session['permissions'] = permissions
            session['role']        = role
            session['login_time']  = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            session.modified = True
            print("✅ SESSION SET:", dict(session))

            # ✅ طباعة للتأكد من نجاح التخزين
            print("✅ تم تسجيل الدخول - SESSION قبل التخزين:", dict(session))

            flash("✅ تم تسجيل الدخول بنجاح", "success")
            return redirect(url_for('index'))

        else:
            flash("⛔ اسم المستخدم أو كلمة المرور غير صحيحة", "danger")
            return redirect(url_for('login'))

    return render_template('login.html')


    hashed_password = generate_password_hash(password)

    # إدخال بيانات المستخدم في جدول users
    c.execute('''
        INSERT INTO users (first_name, last_name, email, username, password, company, facility, role, is_active)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (first_name, last_name, email, username, hashed_password, company, facility, 'admin', 1))
    conn.commit()

    # إعطاء المستخدم جميع الصلاحيات افتراضياً
    c.execute("SELECT id FROM users WHERE username = ?", (username,))
    user_id = c.fetchone()[0]
    full_permissions = [
        'upload_pm', 'upload_asset', 'delete_excel', 'generate_reports',
        'view_edit', 'view_edit_asset', 'view_dashboard', 'can_add_users'
    ]
    for perm in full_permissions:
        c.execute("INSERT INTO user_permissions (user_id, permission) VALUES (?, ?)", (user_id, perm))

    conn.commit()
    conn.close()

    flash("✅ تم إنشاء الحساب بنجاح!", "success")
    return redirect(url_for('index'))
#🔷****
@app.route('/edit_user', methods=['POST'])
def edit_user():
    username = request.form.get('username')
    password = request.form.get('password')
    action = request.form.get('edit_action')

    conn = sqlite3.connect("data.db")
    c = conn.cursor()

    # التحقق من اسم المستخدم وكلمة المرور
    c.execute("SELECT id, password FROM users WHERE username = ?", (username,))
    user = c.fetchone()

    if not user or not check_password_hash(user[1], password):
        flash("❌ اسم المستخدم أو كلمة المرور غير صحيحة.", "danger")
        return redirect(url_for('login'))

    user_id = user[0]

    # تسجيل الدخول
    session.clear()
    session.permanent = True
    app.permanent_session_lifetime = timedelta(hours=6)
    session['user_id'] = user_id
    session['username'] = username
    session['login_time'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # تحميل باقي بيانات الجلسة
    c.execute("SELECT role, company, facility FROM users WHERE id = ?", (user_id,))
    udata = c.fetchone()
    session['role'] = udata[0]
    session['company'] = udata[1]
    session['facility'] = udata[2]

    # تحميل الصلاحيات
    c.execute("SELECT permission FROM user_permissions WHERE user_id = ?", (user_id,))
    session['permissions'] = [row[0] for row in c.fetchall()]
    conn.close()

    flash("✅ تم تسجيل الدخول بنجاح", "success")

    # تحويل المستخدم حسب اختياره
    if action == "permissions":
        return redirect(url_for('manage_users', username=username))
    elif action == "password":
        return redirect(url_for('reset_password'))
    else:
        return redirect(url_for('index'))

#🔷****
#🔷****
#🔷****
### تغيير كلمة مرور المستخدم الحالي من خلال صفحة reset_password

@app.route('/reset_password', methods=['GET', 'POST'])
def reset_password():
    if request.method == 'POST':
        new_password = request.form.get('new_password')
        if not new_password:
            flash("❌ كلمة المرور الجديدة مطلوبة.", "danger")
            return render_template("reset_password.html")

        # استخدام اسم المستخدم الموجود في الجلسة أو "admin" كافتراضي
        username = session.get('username', 'admin')

        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        hashed_pw = generate_password_hash(new_password)
        c.execute("UPDATE users SET password = ? WHERE username = ?", (hashed_pw, username))
        conn.commit()
        conn.close()

        flash("✅ تم تغيير كلمة المرور بنجاح.", "success")
        return redirect(url_for('index'))

    return render_template("reset_password.html")
#🔷****
#🔷****
#🔷****
#🔷****
### إعداد مسارات مجلدات التحميل والصور، وإنشاء مجلد images إذا لم يكن موجود، وتحديد أنواع الملفات المسموح بها

# إعدادات الملفات والصور
BASE_DIR = os.path.dirname(os.path.abspath(__file__))  # المسار الأساسي للملف الحالي
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')      # مجلد التحميل الرئيسي
IMAGES_FOLDER = os.path.join(UPLOAD_FOLDER, 'Images')  # مجلد فرعي للصور
os.makedirs(IMAGES_FOLDER, exist_ok=True)              # إنشاء مجلد الصور إذا لم يكن موجود

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER            # تعيين إعداد التحميل في Flask

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}     # الأنواع المسموحة للصور
#🔷****
#🔷****
#🔷****
#🔷****
### دوال مساعدة متنوعة تُستخدم في عدة أجزاء من التطبيق

# ✅ تنظيف أسماء الملفات من الرموز غير المسموحة لتكون صالحة للتخزين
def clean_filename(name):
    return re.sub(r'[\\/:#🔷****?"<>|#]', '_', name.strip().replace(' ', '_'))

# ✅ التحقق إذا كان الملف من نوع مسموح (صورة)
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# ✅ توليد هاش باستخدام SHA256 لكلمة مرور (لم يتم استخدامه فعليًا لأن `generate_password_hash` من werkzeug يُستخدم)
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# ✅ التحقق مما إذا كان المستخدم لديه صلاحية معينة
def has_permission(user_id, permission):
    conn = sqlite3.connect('data.db')
    c = conn.cursor()
    c.execute("SELECT 1 FROM user_permissions WHERE user_id=? AND permission=?", (user_id, permission))
    result = c.fetchone()
    conn.close()
    return result is not None

# ✅ جلب معلومات آخر ملف credentials.json تم رفعه (الاسم وتاريخ الرفع)
def get_credentials_status():
    with sqlite3.connect("data.db") as conn:
        c = conn.cursor()
        c.execute("SELECT filename, uploaded_at FROM credentials_files ORDER BY id DESC LIMIT 1")
        return c.fetchone()

# ✅ جلب محتوى ملف credentials.json من قاعدة البيانات (لاستخدامه في Google Drive API)
def get_credentials_from_db():
    with sqlite3.connect("data.db") as conn:
        c = conn.cursor()
        c.execute("SELECT content FROM credentials_files ORDER BY id DESC LIMIT 1")
        result = c.fetchone()
        if result:
            return json.loads(result[0])
    return None
#🔷****
#🔷****
#🔷****
#🔷****
### رفع ملف credentials.json وتخزينه داخل قاعدة البيانات لاستخدامه لاحقًا في Google Drive API

from flask import Flask, render_template, request, redirect, url_for, flash
import sqlite3
from datetime import datetime

@app.route('/upload_credentials', methods=['POST'])
def upload_credentials():
    file = request.files.get('credentials_file')

    # ✅ التحقق من أن الملف موجود وصيغته JSON
    if not file or not file.filename.endswith('.json'):
        flash("❌ الرجاء رفع ملف بصيغة JSON صحيحة.", "danger")
        return redirect(url_for('index'))

    # ✅ قراءة محتوى الملف
    content = file.read()

    # ✅ الاتصال بقاعدة البيانات وتحديث البيانات
    with sqlite3.connect("data.db") as conn:
        c = conn.cursor()

        # إنشاء الجدول إذا لم يكن موجوداً (احتياطياً)
        c.execute('''
            CREATE TABLE IF NOT EXISTS credentials_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                content BLOB,
                uploaded_at TEXT
            )
        ''')

        # ✅ حذف أي ملفات اعتماد سابقة
        c.execute("DELETE FROM credentials_files")

        # ✅ إدخال الملف الجديد بالتاريخ
        uploaded_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        c.execute(
            "INSERT INTO credentials_files (filename, content, uploaded_at) VALUES (?, ?, ?)",
            (file.filename, content, uploaded_at)
        )

        conn.commit()

    # ✅ عرض رسالة نجاح
    flash("✅ تم رفع ملف الاعتماد بنجاح.", "success")
    return redirect(url_for('index'))

#🔷****
#🔷****
#🔷****
#🔷****
### جلب سجل معين (صف) من ملف Excel مرفوع (PM أو Asset) بناءً على نوع الملف واسم الشيت ورقم الصف

@app.route('/get_record')
def get_record():
    try:
        data_type = request.args.get("type", "").lower()
        if not data_type:
            return jsonify({"error": "Missing type parameter"}), 400

        sheet_name = request.args.get("sheet")
        row_index = int(request.args.get("row"))
        header_row = 0 if data_type == 'pm' else 1

        sheets = load_all_sheets_from_db(data_type, header_row)
        df = sheets.get(sheet_name)
        if df is None:
            return jsonify({"error": "Sheet not found"}), 404

        row = df.iloc[row_index].fillna('').to_dict()

        return jsonify({
            "data": row,
            "sheet": sheet_name,
            "row": row_index,
            "type": data_type
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
#🔷****
@app.route('/check_if_already_saved')
def check_if_already_saved():
    sheet = request.args.get("sheet")
    row = request.args.get("row", type=int)
    entry_type = request.args.get("type")

    if not sheet or row is None or not entry_type:
        return jsonify({"error": "❌ Missing parameters"}), 400

    try:
        with sqlite3.connect("data.db") as conn:
            c = conn.cursor()
            c.execute("""
                SELECT 1 FROM maintenance_records
                WHERE sheet_name = ? AND row_index = ? AND LOWER(type) = LOWER(?)
                LIMIT 1
            """, (sheet, row, entry_type))
            found = c.fetchone()

        return jsonify({"already_saved": bool(found)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
#🔷****
#🔷****
#🔷****
### حفظ سجل PM بعد التعديل، وإنشاء الملفات المرتبطة (Word, Excel, ZIP)، ورفع ZIP إلى Google Drive

@app.route('/save_edit', methods=['POST'])
def save_edit():
    return handle_save(request, entry_type='PM')
#🔷****
#🔷****
#🔷****
#🔷****
### حفظ سجل Asset بعد التعديل (نفس الوظيفة الأساسية لكن لـ Asset بدل PM)

@app.route('/save_asset_edit', methods=['POST'])
def save_asset_edit():
    return handle_save(request, entry_type='Asset')
#🔷****
#🔷****
@app.route('/disable_user/<username>', methods=['POST'])
def disable_user(username):
    conn = sqlite3.connect("data.db")
    c = conn.cursor()
    c.execute("UPDATE users SET is_active = 0 WHERE username = ?", (username,))
    conn.commit()
    conn.close()
    flash(f"⛔ تم تعطيل المستخدم {username}", "warning")
    return redirect(url_for('manage_users'))

#🔷****
@app.route('/delete_user/<username>', methods=['POST'])
def delete_user(username):
    conn = sqlite3.connect("data.db")
    c = conn.cursor()
    c.execute("DELETE FROM users WHERE username = ?", (username,))
    conn.commit()
    conn.close()
    flash(f"✅ تم حذف المستخدم {username}", "success")
    return redirect(url_for('manage_users'))

#🔷****
### دالة handle_save تقوم بـ:
### - استقبال البيانات والصور
### - حفظها في قاعدة البيانات
### - إنشاء ملف Excel وWord
### - حفظ الصور في مجلد
### - ضغط كل شيء في ملف ZIP
### - رفع الملف إلى Google Drive
### - حفظ ZIP في قاعدة البيانات
### - إرجاع نتائج العملية للمستخدم

def handle_save(req, entry_type):
    try:
        sheet_name = req.form.get('sheet_name')
        row_index = int(req.form.get('row_index'))
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # ✅ جلب الحقول التي تبدأ بـ field_
        data = {key[6:]: req.form.get(key) for key in req.form if key.startswith('field_')}

        # ✅ توليد ملاحظة تلقائية من الجلسة إذا لم تُكتب يدويًا
        notes_text = req.form.get('notes_text', '').strip()
        if not notes_text:
            notes_text = f"{session.get('username', '')} - {session.get('company', '')} - {session.get('facility', '')}"

        # ✅ تعبئة الحقول تلقائيًا
        data['Notes'] = notes_text
        data['Remarks'] = notes_text
        data['Attended by'] = f"{session.get('username', '')} - {session.get('company', '')} - {session.get('facility', '')}"

        # ✅ تعيين الأعمدة المهمة حسب نوع الإدخال
        if entry_type == 'PM':
            search_columns = [
                "Work Order", "Asset Number", "Description", "Priority",
                "Department", "Planner Group", "Work Center", "Reported By"
            ]
        elif entry_type == 'Asset':
            search_columns = [
                "Asset Number", "Asset Description", "Category", "Location",
                "Cost Center", "Manufacturer", "Model", "Serial Number"
            ]
        else:
            search_columns = []


        # ✅ تحميل الصور
        def save_files(files_list):
            saved_files = []
            for file in files_list:
                if file and allowed_file(file.filename):
                    filename = secure_filename(f"{datetime.now().strftime('%Y%m%d%H%M%S%f')}_{file.filename}")
                    save_path = os.path.join(IMAGES_FOLDER, filename)
                    file.save(save_path)
                    saved_files.append(filename)
            return saved_files

        before_images = save_files(req.files.getlist('Before Maintenance')) if entry_type == 'PM' else []
        after_images = save_files(req.files.getlist('After Maintenance')) if entry_type == 'PM' else []
        report_images = save_files(req.files.getlist('Maintenance Report')) if entry_type == 'PM' else []
        cm_images = save_files(req.files.getlist('CM Images')) if entry_type in ['CM', 'Asset'] else []
        spare_parts_images = save_files(req.files.getlist('Spare Parts Images')) if entry_type == 'Asset' else []
        notes_images = save_files(req.files.getlist('notes_images'))

        # ✅ حفظ البيانات في قاعدة البيانات
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute('''
            INSERT INTO maintenance_records (
                type, sheet_name, row_index, data, timestamp,
                before_images, after_images, report_images,
                cm_images, notes_text, notes_images
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            entry_type, sheet_name, row_index, json.dumps(data, ensure_ascii=False), timestamp,
            ','.join(before_images), ','.join(after_images), ','.join(report_images),
            ','.join(cm_images), notes_text, ','.join(notes_images)
        ))
        conn.commit()
        conn.close()

        # ✅ تجهيز ملف Excel
        data['Notes'] = notes_text
        df = pd.DataFrame([data])
        df = df[[col for col in search_columns if col in df.columns] + ['Notes']]

        wo = clean_filename(data.get('Work Order') or data.get('Asset') or data.get('Ticket No') or 'UnknownWO')
        location = clean_filename(data.get('Location') or 'UnknownLoc')
        desc = clean_filename(data.get('Description') or 'NoDesc')[:30]
        file_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{wo}_{file_timestamp}_{location}_{desc}"

        output_folder = os.path.join(UPLOAD_FOLDER, base_filename)
        os.makedirs(output_folder, exist_ok=True)

        excel_file = os.path.join(output_folder, f"{base_filename}.xlsx")
        df.to_excel(excel_file, sheet_name=f'{entry_type}_Records', index=False)

        append_to_daily_excel(entry_type, data, timestamp)

        # ✅ تجهيز ملف Word
        word_path = os.path.join(output_folder, f"{base_filename}.docx")
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        logo_path = os.path.join(BASE_DIR, 'Metco-logo.png')
        if os.path.exists(logo_path):
            header_section = doc.sections[0].header
            paragraph = header_section.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_heading(f"Work Order: {data.get('Work Order', data.get('Asset', ''))}", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for key, val in data.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = str(val)
            for cell in row:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        def add_images_table(images_list, label):
            if not images_list:
                return
            doc.add_paragraph(label, style='Heading 2')
            img_table = doc.add_table(rows=0, cols=3)
            img_table.style = 'Table Grid'
            for i in range(0, len(images_list), 3):
                row_cells = img_table.add_row().cells
                for j, img in enumerate(images_list[i:i + 3]):
                    img_path = os.path.join(IMAGES_FOLDER, img)
                    if os.path.exists(img_path):
                        run = row_cells[j].paragraphs[0].add_run()
                        run.add_picture(img_path, width=Inches(2))
                        row_cells[j].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        add_images_table(report_images, 'Report')
        add_images_table(before_images, 'Before')
        add_images_table(after_images, 'After')
        add_images_table(cm_images, 'CM Images')
        add_images_table(spare_parts_images, 'Spare Parts')
        if notes_text:
            doc.add_paragraph('Notes:', style='Heading 2')
            doc.add_paragraph(notes_text)
        add_images_table(notes_images, 'Notes')

        doc.save(word_path)

        for img in before_images + after_images + report_images + cm_images + notes_images + spare_parts_images:
            shutil.move(os.path.join(IMAGES_FOLDER, img), os.path.join(output_folder, img))

        zip_filename = f"{base_filename}.zip"
        zip_path = os.path.join(UPLOAD_FOLDER, zip_filename)
        shutil.make_archive(zip_path.replace('.zip', ''), 'zip', output_folder)

        upload_to_drive(zip_path, zip_filename)
        save_zip_to_db(zip_filename, zip_path)

        timers = []
        for field in data:
            if "Traget finish" in field or "Target Finish" in field:
                try:
                    finish_time = parser.parse(data[field])
                    timers.append({
                        "field": field,
                        "expiresAt": int(finish_time.timestamp() * 1000)
                    })
                except Exception as parse_error:
                    print("⛔ Error parsing timer field:", field, parse_error)

        response_data = {
            "message": "Saved",
            "zip_file": zip_filename
        }

        if timers:
            response_data["timers"] = timers
            response_data["ticket"] = data.get("Ticket No", "Unknown")
            response_data["serial"] = data.get("Serial Number") or data.get("Sr. No") or data.get("Asset", "Unknown")

        return render_template("save_result.html",
                               zip_file=zip_filename,
                               entry_type=entry_type,
                               sheet_name=sheet_name,
                               row_index=row_index)

    except Exception as e:
        print("ERROR in handle_save:", e)
        return jsonify({"error": str(e)}), 500

#🔷****
#🔷****
#🔷****
#🔷****
### حفظ السجل الجديد في ملف Excel يومي داخل قاعدة البيانات (Daily PM TKT أو Daily CM TKT)

def append_to_daily_excel(entry_type, new_data, timestamp):
    filename = "Daily PM TKT.xlsx" if entry_type == "PM" else "Daily CM TKT.xlsx"
    sheet_name = "PM Records" if entry_type == "PM" else "CM Records"

    df_new = pd.DataFrame([new_data])
    df_new['timestamp'] = timestamp

    with sqlite3.connect("data.db") as conn:
        c = conn.cursor()

        # محاولة جلب آخر نسخة من الملف اليومي
        c.execute("SELECT content FROM daily_excel_files WHERE filename = ? ORDER BY id DESC LIMIT 1", (filename,))
        result = c.fetchone()

        if result:
            existing_data = io.BytesIO(result[0])
            excel_file = pd.ExcelFile(existing_data)
            sheets = {s: pd.read_excel(excel_file, sheet_name=s) for s in excel_file.sheet_names}
        else:
            sheets = {}

        if sheet_name in sheets:
            df_existing = sheets[sheet_name]
            df_combined = pd.concat([df_existing, df_new], ignore_index=True)
        else:
            df_combined = df_new

        sheets[sheet_name] = df_combined

        # حفظ النسخة المحدثة في قاعدة البيانات
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for sheet, df in sheets.items():
                df.to_excel(writer, sheet_name=sheet, index=False)

        output.seek(0)

        c.execute("DELETE FROM daily_excel_files WHERE filename = ?", (filename,))
        c.execute("INSERT INTO daily_excel_files (filename, content, updated_at) VALUES (?, ?, ?)",
                  (filename, output.read(), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
        conn.commit()
#🔷****
#🔷****
#🔷****
#🔷****
### حذف ملف Excel الذي تم توليده (PM TKT أو CM TKT) من جدول maintenance_excel في قاعدة البيانات

@app.route('/delete_generated_excel/<entry_type>', methods=['POST'])
def delete_generated_excel(entry_type):
    filename = "PM TKT.xlsx" if entry_type == 'pm' else "CM TKT.xlsx"
    try:
        with sqlite3.connect("data.db") as conn:
            c = conn.cursor()
            c.execute("DELETE FROM maintenance_excel WHERE filename = ?", (filename,))
            conn.commit()
        return redirect(url_for('index'))
    except Exception as e:
        return f"Error deleting file: {e}", 500
#🔷****
#🔷****
#🔷****
#🔷****
### توليد ملف Excel يحتوي على السجلات من جدول maintenance_records بين تاريخين محددين، بحسب النوع (PM أو Asset)، وحفظه في قاعدة البيانات

@app.route('/generate_excel', methods=['POST'])
def generate_excel():
    try:
        entry_type = request.form.get('type')                    # نوع الإدخال: 'pm' أو 'asset'
        start_date = request.form.get('start_date')              # تاريخ البداية
        end_date = request.form.get('end_date')                  # تاريخ النهاية

        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT data, timestamp, notes_text FROM maintenance_records WHERE LOWER(type) = ?", (entry_type.lower(),))
        records = c.fetchall()
        conn.close()

        filtered = []
        for data_json, timestamp, notes_text in records:
            try:
                ts_date = datetime.strptime(timestamp, "%Y-%m-%d %H:%M:%S").date()
                start = datetime.strptime(start_date, "%Y-%m-%d").date()
                end = datetime.strptime(end_date, "%Y-%m-%d").date()
            except:
                continue

            if start <= ts_date <= end:
                row = json.loads(data_json)
                row['timestamp'] = timestamp
                row['Notes'] = notes_text or ''
                filtered.append(row)

        if not filtered:
            return render_template("index.html", message="⚠️ لا توجد بيانات بين التاريخين المحددين")

        df = pd.DataFrame(filtered)

        # ترتيب الأعمدة إذا كان النوع PM
        if entry_type == 'pm':
            ordered_columns = [
                "Work Order", "PM", "Job Plan", "Parent WO", "Description", "Location", "Asset",
                "MMS #", "QR CODE", "Route", "Work Type", "Workshop", "Target Start", "Target Finish", "METCO COMMENT"
            ]
            df = df[[col for col in ordered_columns if col in df.columns] + [col for col in df.columns if col not in ordered_columns]]

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            sheet = "PM List Results" if entry_type == 'pm' else "Asset List Results"
            df.to_excel(writer, index=False, sheet_name=sheet)
        output.seek(0)

        filename = "PM TKT.xlsx" if entry_type == 'pm' else "CM TKT.xlsx"

        with sqlite3.connect("data.db") as conn:
            c = conn.cursor()
            c.execute("DELETE FROM maintenance_excel WHERE filename = ?", (filename,))
            c.execute("INSERT INTO maintenance_excel (filename, content, updated_at) VALUES (?, ?, ?)",
                      (filename, output.read(), datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            conn.commit()

        output.seek(0)
        return send_file(output, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except Exception as e:
        return f"Error generating Excel: {e}", 500
#🔷****
#🔷****
#🔷****
#🔷****
### حذف ملف ZIP واحد من قاعدة البيانات حسب الاسم

@app.route('/delete_zip', methods=['POST'])
def delete_zip():
    try:
        filename = request.json.get('filename')
        if not filename:
            return jsonify({"error": "Filename is required"}), 400

        conn = sqlite3.connect("data.db")
        c = conn.cursor()
        c.execute("DELETE FROM zip_files WHERE filename = ?", (filename,))
        conn.commit()
        conn.close()
        return jsonify({"message": f"{filename} deleted successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
#🔷****
#🔷****
#🔷****
#🔷****
### حذف مجموعة من ملفات ZIP دفعة واحدة من قاعدة البيانات


@app.route('/delete_zip_batch', methods=['POST'])
def delete_zip_batch():
    try:
        files = request.json.get('files', [])
        print("🟡 Received for batch delete:", files)  # ← أضف هذه السطر

        if not files:
            return jsonify({"error": "No files provided"}), 400

        ...
        conn = sqlite3.connect("data.db")
        c = conn.cursor()
        c.executemany("DELETE FROM zip_files WHERE filename = ?", [(f,) for f in files])
        conn.commit()
        conn.close()
        return jsonify({"message": "Files deleted successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
#🔷****
#🔷****
#🔷****
#🔷****
### رفع ملف ZIP إلى Google Drive باستخدام credentials المخزنة في قاعدة البيانات

def upload_to_drive(zip_path, filename):
    try:
        SCOPES = ['https://www.googleapis.com/auth/drive.file']  # صلاحيات الوصول إلى Google Drive
        creds_info = get_credentials_from_db()
        if not creds_info:
            print("❌ No credentials found in DB")
            return

        # إنشاء Credentials من محتوى JSON المحفوظ
        creds = service_account.Credentials.from_service_account_info(creds_info, scopes=SCOPES)
        service = build('drive', 'v3', credentials=creds)

        folder_id = '1SolVjxUU0iZ7YRgmgt424_JMtkPU6CUG'  # معرف المجلد في Google Drive
        file_metadata = {'name': filename, 'parents': [folder_id]}
        media = MediaFileUpload(zip_path, mimetype='application/zip')

        # تنفيذ عملية الرفع
        service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    except Exception as e:
        print("Google Drive upload failed:", e)
#🔷****
#🔷****
#🔷****
#🔷****
### إعادة توجيه المستخدم من الجذر `/` إلى صفحة تسجيل الدخول
@app.route('/')
def index():
    # ✅ التحقق من تسجيل الدخول
    if 'username' not in session:
        return redirect(url_for('login'))

    # ✅ جلب بيانات ملف credentials من قاعدة البيانات
    with sqlite3.connect("data.db") as conn:
        c = conn.cursor()
        c.execute("SELECT filename, uploaded_at FROM credentials_files ORDER BY uploaded_at DESC LIMIT 1")
        row = c.fetchone()
        credentials_status = row if row else None

    # ✅ عرض الصفحة الرئيسية مع تمرير اسم المستخدم وبيانات الاعتماد
    return render_template("index.html",
                           session_username=session.get('username'),
                           session_company=session.get('company'),
                           session_facility=session.get('facility'),
                           session_login_time=session.get('login_time'),
                           credentials_status=credentials_status)
#🔷****
#🔷****
#🔷****
#🔷****
### تحميل ملف ZIP من قاعدة البيانات باستخدام الـ ID الخاص به

@app.route('/download_zip/<int:zip_id>')
def download_zip(zip_id):
    try:
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT filename, content FROM zip_files WHERE id=?", (zip_id,))
        row = c.fetchone()
        conn.close()
        if row:
            filename, content = row
            return send_file(io.BytesIO(content), as_attachment=True, download_name=filename, mimetype='application/zip')
        else:
            return "ZIP not found", 404
    except Exception as e:
        return f"Error: {e}", 500
#🔷****
#🔷****
#🔷****
#🔷****
### تحميل ملف ZIP باستخدام اسمه كـ parameter في المسار

@app.route('/download_zip/<path:filename>')
def download_zip_filename(filename):
    try:
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT content FROM zip_files WHERE filename = ?", (filename,))
        row = c.fetchone()
        conn.close()
        if row:
            content = row[0]
            return send_file(io.BytesIO(content), as_attachment=True, download_name=filename, mimetype='application/zip')
        else:
            return "ZIP file not found", 404
    except Exception as e:
        return f"Error: {e}", 500
#🔷****
#🔷****
#🔷****
#🔷****
### تحميل ملف ZIP باستخدام اسمه مباشرة (مفيد للاستخدام مع JavaScript)

@app.route('/download_zip_by_name/<path:filename>')
def download_zip_by_name(filename):
    try:
        conn = sqlite3.connect('data.db')
        c = conn.cursor()
        c.execute("SELECT content FROM zip_files WHERE filename = ?", (filename,))
        row = c.fetchone()
        conn.close()
        if row:
            content = row[0]
            return send_file(io.BytesIO(content), as_attachment=True, download_name=filename, mimetype='application/zip')
        else:
            return "ZIP file not found", 404
    except Exception as e:
        return f"Error: {e}", 500
#🔷****
#🔷****
#🔷****
#🔷****
### البحث في قاعدة البيانات عن ملفات ZIP تحتوي على الكلمة المفتاحية في الاسم

@app.route('/search_zip', methods=['POST'])
def search_zip():
    try:
        query = request.json.get('query', '').strip().lower()
        conn = sqlite3.connect('data.db')
        cursor = conn.cursor()
        cursor.execute("SELECT filename FROM zip_files WHERE LOWER(filename) LIKE ?", (f"%{query}%",))
        files = [row[0] for row in cursor.fetchall()]
        conn.close()
        return jsonify(files)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
#🔷****
#🔷****
#🔷****
#🔷****
### حفظ ملف Excel داخل قاعدة البيانات بحسب نوعه (pm أو asset)

def save_uploaded_excel(file, file_type):
    data = file.read()
    with sqlite3.connect('data.db') as conn:
        c = conn.cursor()
        c.execute("DELETE FROM uploaded_excels WHERE type = ?", (file_type,))
        c.execute("INSERT INTO uploaded_excels (type, filename, data) VALUES (?, ?, ?)",
                  (file_type, file.filename, data))
        conn.commit()
#🔷****
#🔷****
#🔷****
#🔷****
### استرجاع أحدث ملف Excel مخزن في قاعدة البيانات حسب النوع

def get_uploaded_excel(file_type):
    with sqlite3.connect('data.db') as conn:
        c = conn.cursor()
        c.execute("SELECT filename, data, created_at FROM uploaded_excels WHERE type = ? ORDER BY id DESC LIMIT 1", (file_type,))
        return c.fetchone()
#🔷****
#🔷****
#🔷****
#🔷****
### حذف ملف Excel من قاعدة البيانات حسب نوعه (pm أو asset)

def delete_uploaded_excel(file_type):
    with sqlite3.connect('data.db') as conn:
        c = conn.cursor()
        c.execute("DELETE FROM uploaded_excels WHERE type = ?", (file_type,))
        conn.commit()
#🔷****
#🔷****
#🔷****
#🔷****
### تحميل جميع الشيتات من ملف Excel المخزن في قاعدة البيانات وتحويلها إلى DataFrames

def load_all_sheets_from_db(file_type, header_row):
    result = get_uploaded_excel(file_type)
    if not result:
        return {}
    _, data, _ = result
    xls = pd.ExcelFile(io.BytesIO(data))
    return {sheet: pd.read_excel(xls, sheet_name=sheet, header=header_row) for sheet in xls.sheet_names}
#🔷****
#🔷****
#🔷****
#🔷****
### مسار رفع ملف Excel جديد وتخزينه في قاعدة البيانات

@app.route('/upload_excel', methods=['POST'])
def upload_excel():
    file = request.files['file']
    file_type = request.form['type']
    if file and file_type in ['pm', 'asset']:
        save_uploaded_excel(file, file_type)
    return redirect(url_for('index'))
#🔷****
#🔷****
#🔷****
#🔷****
### مسار حذف ملف Excel من قاعدة البيانات بناءً على نوعه

@app.route('/delete_excel/<file_type>', methods=['POST'])
def delete_excel(file_type):
    if file_type in ['pm', 'asset']:
        delete_uploaded_excel(file_type)
    return redirect(url_for('index'))
#🔷****
#🔷****
#🔷****
#🔷****
### البحث داخل ملفات Excel المرفوعة عن كلمة مفتاحية ضمن كل الصفوف في جميع الشيتات

@app.route('/search', methods=['GET'])
def search():
    try:
        keyword = request.args.get('keyword', '').strip().lower()
        source = request.args.get('source', '').strip().lower()
        results = []

        if source not in ['pm', 'asset'] or not keyword:
            return jsonify({'results': results})

        header_row = 0 if source == 'pm' else 1
        sheets = load_all_sheets_from_db(source, header_row)

        for sheet_name, df in sheets.items():
            df = df.fillna('')
            for idx, row in df.iterrows():
                if any(keyword in str(val).lower() for val in row):
                    result = row.to_dict()
                    result['SheetName'] = sheet_name
                    result['RowIndex'] = int(idx)
                    results.append(result)

        return jsonify({'results': results})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
#🔷****
#🔷****
#🔷****
#🔷****
### عرض صفحة تعديل PM – مع التحقق من صلاحية الوصول بناءً على الجلسة أو اسم المستخدم

@app.route('/edit')
def edit_page():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    # ✅ السماح المطلق للمستخدم Ar
    if session.get('username') != 'Ar' and not has_permission(session['user_id'], 'view_edit'):
       return "⛘️ You do not have permission to access this page", 403

    return render_template('edit_page.html',
                           username=session.get('username'),
                           email=session.get('email'))
#🔷****
#🔷****
#🔷****
#🔷****

### عرض صفحة تعديل Asset – تمرير اسم الشيت ورقم الصف كمُعاملات

@app.route('/edit_asset')
def edit_asset_page():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    if not has_permission(session['user_id'], 'view_edit_asset'):
        return "⛘️ You do not have permission to access this page", 403

    sheet = request.args.get('sheet')
    row = request.args.get('row')
    return render_template('edit_asset.html',
                           sheet_name=sheet,
                           row_index=row,
                           username=session.get('username'),
                           email=session.get('email'))
#🔷****

#🔷****
@app.route('/permissions_admin', methods=['GET'])
def permissions_admin():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("⛔ You do not have permission to access this page.", "danger")
        return redirect(url_for('index'))

    all_permissions = [
        'upload_pm', 'upload_asset', 'delete_excel',
        'generate_reports', 'edit_records',
        'view_zip', 'add_users', 'view_edit_asset', 'view_edit'
    ]

    conn = sqlite3.connect('data.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()

    c.execute("SELECT id, username FROM users WHERE is_active = 1")
    users = []
    for row in c.fetchall():
        c.execute("SELECT permission FROM user_permissions WHERE user_id = ?", (row['id'],))
        perms = [r['permission'] for r in c.fetchall()]
        users.append({
            'id': row['id'],
            'username': row['username'],
            'permissions': perms
        })

    conn.close()
    return render_template(
        'permissions_admin.html',
        users=users,
        all_permissions=all_permissions
    )

#🔷****
@app.route('/update_all_permissions', methods=['POST'])
def update_all_permissions():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash("⛔ لا تملك صلاحية الدخول.", "danger")
        return redirect(url_for('index'))

    conn = sqlite3.connect('data.db')
    c = conn.cursor()

    c.execute("SELECT id FROM users WHERE is_active = 1")
    user_ids = [row[0] for row in c.fetchall()]

    for user_id in user_ids:
        c.execute("DELETE FROM user_permissions WHERE user_id = ?", (user_id,))
        new_perms = request.form.getlist(f"perms_{user_id}[]")
        for p in new_perms:
            c.execute("INSERT INTO user_permissions (user_id, permission) VALUES (?, ?)", (user_id, p))
        manual_perm = request.form.get(f"new_perm_{user_id}", "").strip().lower()
        if manual_perm:
            c.execute("INSERT INTO user_permissions (user_id, permission) VALUES (?, ?)", (user_id, manual_perm))

    conn.commit()
    conn.close()
    flash("✅ تم تحديث جميع الصلاحيات بنجاح.", "success")
    return redirect(url_for('permissions_admin'))

#🔷****
### عرض أي صورة أو ملف تم رفعه إلى مجلد /uploads عند الطلب

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)
#🔷****
#🔷****
#🔷****
#🔷****
### جلب اسم وتاريخ آخر ملف Excel مرفوع (PM أو Asset) لعرضه في الواجهة

@app.route('/get_uploaded_excel_info/<file_type>')
def get_uploaded_excel_info(file_type):
    if file_type not in ['pm', 'asset']:
        return jsonify({"error": "Invalid file type"}), 400

    with sqlite3.connect("data.db") as conn:
        c = conn.cursor()
        c.execute("SELECT filename, created_at FROM uploaded_excels WHERE type = ? ORDER BY id DESC LIMIT 1", (file_type,))
        row = c.fetchone()
        if row:
            return jsonify({"filename": row[0], "created_at": row[1]})
        else:
            return jsonify({})
#🔷****
#🔷****
#🔷****
#🔷****
### حفظ ملف ZIP داخل قاعدة البيانات بعد إنشائه

def save_zip_to_db(filename, zip_path):
    with open(zip_path, 'rb') as f:
        content = f.read()
    conn = sqlite3.connect('data.db')
    c = conn.cursor()
    c.execute("INSERT INTO zip_files (filename, content) VALUES (?, ?)", (filename, content))
    conn.commit()
    conn.close()
#🔷****
#🔷****
#🔷****
#🔷****
#🔷****
@app.route('/logout', methods=['POST'])
def logout():
    session.clear()
    return render_template("login.html")
#🔷****
@app.route('/chat')
def chat_page():
    if 'username' not in session:
        return redirect(url_for('login'))

    conn = sqlite3.connect('data.db')
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT username, message FROM chat_messages ORDER BY timestamp ASC LIMIT 50")
    messages = c.fetchall()
    conn.close()

    return render_template('chat.html', username=session['username'], messages=messages)

@socketio.on('send_message')
def handle_message(data):
    username = data.get('user')
    message = data.get('message')
    image = data.get('image')
    audio = data.get('audio')

    if not username:
        print("لا يوجد اسم مستخدم")
        return

    content = None
    if message:
        content = '[TEXT] ' + message
    elif image:
        content = '[IMAGE] ' + image
    elif audio:
        content = '[AUDIO] ' + audio

    if not content:
        print("⚠️ رسالة فارغة - لا يتم الحفظ")
        return

    try:
        with sqlite3.connect('data.db') as conn:
            c = conn.cursor()
            c.execute("INSERT INTO chat_messages (username, message) VALUES (?, ?)", (username, content))
            conn.commit()
    except Exception as e:
        print("❌ Error saving to chat_messages:", e)

    # حفظ في json history file
    save_message(username, {'text': message} if message else {'image': image} if image else {'audio': audio})

    # إرسال للجميع
    emit('receive_message', data, broadcast=True)


@app.route('/chat_history')
def chat_history():
    since = datetime.now() - timedelta(days=7)
    messages = [
        m for m in load_history()
        if datetime.fromisoformat(m['timestamp']) >= since
    ]
    return jsonify(messages)

#🔷****
#🔷****
#🔷****
#🔷****
### تشغيل تطبيق Flask عند تنفيذ الملف مباشرة، على المنفذ 5000 وعلى كل العناوين (0.0.0.0)

if __name__ == '__main__':
    socketio.run(app, host='0.0.0.0', port=5000, debug=True, use_reloader=False)
#🔷****
#🔷****
#🔷****
#🔷****